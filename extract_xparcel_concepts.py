"""
Extract text from Xparcel Core Concepts Explained.docx
"""

import docx

# Load the document
doc_path = r"c:\Users\BrettWalker\Downloads\Xparcel Core Concepts Explained.docx"
doc = docx.Document(doc_path)

print("=" * 80)
print("XPARCEL CORE CONCEPTS EXPLAINED")
print("=" * 80)
print()

# Extract all text
full_text = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        full_text.append(paragraph.text)
        print(paragraph.text)
        print()

# Also check for tables
if doc.tables:
    print("\n" + "=" * 80)
    print("TABLES FOUND IN DOCUMENT")
    print("=" * 80)
    print()

    for i, table in enumerate(doc.tables, 1):
        print(f"Table {i}:")
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            print(" | ".join(row_text))
        print()

print("=" * 80)
print(f"Total paragraphs extracted: {len(full_text)}")
print("=" * 80)
