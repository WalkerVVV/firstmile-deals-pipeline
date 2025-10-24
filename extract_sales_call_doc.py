"""
Extract text from FirstMile & Xparcel_Sales Call.docx and save to markdown
"""

import docx

# Load the document
doc_path = r"c:\Users\BrettWalker\Downloads\FirstMile & Xparcel_Sales Call.docx"
doc = docx.Document(doc_path)

# Open output file with UTF-8 encoding
output_path = "FirstMile_Xparcel_Sales_Call.md"

with open(output_path, 'w', encoding='utf-8') as f:
    f.write("# FirstMile & Xparcel Sales Call\n\n")
    f.write("**Source:** FirstMile Internal Documentation\n\n")
    f.write("---\n\n")

    # Extract all text
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            f.write(paragraph.text + "\n\n")

    # Also check for tables
    if doc.tables:
        f.write("\n---\n\n")
        f.write("## Tables from Document\n\n")

        for i, table in enumerate(doc.tables, 1):
            f.write(f"### Table {i}\n\n")
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                f.write(" | ".join(row_text) + "\n")
            f.write("\n")

print(f"Content extracted successfully to: {output_path}")
print("File saved with UTF-8 encoding to preserve all characters.")
